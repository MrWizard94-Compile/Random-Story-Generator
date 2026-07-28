import requests
import json
import re
from typing import Iterator, List, Tuple, Optional, Dict, Any
from datetime import datetime
import os
import logging
from banned_content import PROMPT_NAME_BAN, validate_story
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Application metadata
__version__ = "1.0.0"
__author__ = "Rob"

# Constants
DEFAULT_WORD_COUNT = 500
DEFAULT_MAX_CHARS = 280
RAPID_MODE_WORD_COUNT = 300
THREAD_POST_INTERVAL_MINUTES = 10
SPINBOX_MIN_WORDS = 100
SPINBOX_MAX_WORDS = 2000


class StoryGenerator:
    """Handles story generation with Ollama models."""

    OLLAMA_API_URL = "http://localhost:11434/api/generate"
    CHAT_API_URL   = "http://localhost:11434/api/chat"
    MODELS_API_URL = "http://localhost:11434/api/tags"

    @staticmethod
    def check_ollama_running() -> bool:
        """Check if Ollama service is running."""
        try:
            response = requests.get(StoryGenerator.MODELS_API_URL, timeout=5)
            return response.status_code == 200
        except Exception:
            return False

    @staticmethod
    def check_ollama_status() -> dict:
        """Check Ollama service status and available models."""
        try:
            response = requests.get(StoryGenerator.MODELS_API_URL, timeout=5)
            if response.status_code == 200:
                try:
                    data = response.json()
                    if not isinstance(data, dict):
                        return {"status": "error", "message": "Invalid API response format"}
                    models = data.get("models", [])
                    if not isinstance(models, list):
                        return {"status": "error", "message": "Invalid models data in API response"}
                    return {
                        "status": "running",
                        "models": [model["name"] for model in models if isinstance(model, dict) and "name" in model],
                        "count": len(models)
                    }
                except json.JSONDecodeError:
                    return {"status": "error", "message": "Invalid JSON response"}
            else:
                return {"status": "error", "message": f"HTTP {response.status_code}"}
        except requests.exceptions.ConnectionError:
            return {"status": "not_running", "message": "Cannot connect to Ollama"}
        except requests.exceptions.Timeout:
            return {"status": "timeout", "message": "Ollama response timeout"}
        except Exception as e:
            return {"status": "error", "message": str(e)}

    @staticmethod
    def get_available_models() -> List[str]:
        """Fetch list of available models from Ollama."""
        try:
            response = requests.get(StoryGenerator.MODELS_API_URL, timeout=10)
            response.raise_for_status()
            data = response.json()
            return [model["name"] for model in data.get("models", [])]
        except requests.exceptions.RequestException as e:
            raise Exception(f"Failed to fetch models: {e}")

    @staticmethod
    def _get_template_text(template: Optional[str]) -> str:
        """Return template-specific instructions."""
        if not template or template.strip() == "None":
            return ""
        template = template.strip().lower()
        templates = {
            "hero's journey": (
                "Structure the story as a classic Hero's Journey: "
                "Ordinary world -> Call to adventure -> Refusal -> Mentor -> Trials -> "
                "Ordeal -> Reward -> The road back -> Resurrection -> Return with elixir. "
                "Emphasize the protagonist's emotional growth, challenge, and transformation."
            ),
            "three-act structure": (
                "Divide the story into three acts: Setup, Confrontation, Resolution. "
                "Act 1 introduces characters and conflict, Act 2 escalates stakes, Act 3 resolves the climax. "
                "Keep clear turning points and character decisions."
            ),
            "five-act structure": (
                "Use five acts (Exposition, Rising Action, Climax, Falling Action, Denouement). "
                "Show cause-and-effect escalation, a strong midpoint, and a satisfying conclusion."
            ),
            "character arc": (
                "Focus on a compelling character arc: beginning flaw, development through conflict, "
                "moment of decision, and final personal change. "
                "Highlight internal conflict and emotional resolution."
            ),
            "mystery detective": (
                "Write it as a mystery/detective structure: crime discovery, investigation, clues, "
                "red herrings, reveal, and logical deduction. Keep suspense and reveal at the climax."
            )
        }
        return templates.get(template, "")

    @staticmethod
    def _build_prompt(genre: Optional[str] = None, tone: Optional[str] = None,
                     word_count: int = DEFAULT_WORD_COUNT, custom_prompt: Optional[str] = None,
                     template: Optional[str] = None) -> str:
        """Build the prompt for story generation."""
        if custom_prompt:
            return custom_prompt + " " + PROMPT_NAME_BAN

        genre_text = f"in the {genre} genre" if genre else ""
        tone_text = f"with a {tone} tone" if tone else ""
        template_text = StoryGenerator._get_template_text(template)

        base = (
            f"Write an original, engaging short story {genre_text} {tone_text}. "
            f"Target length: approximately {word_count} words. "
            f"Include: clear title, characters, setting, plot, conflict, and resolution. "
            f"Ensure proper grammar, punctuation, originality, and no copyrighted material. "
            f"Make it suitable for a general audience. Be creative and imaginative."
        )

        if template_text:
            base += " " + template_text

        # Injected at inference time — system prompt cannot reliably override
        # deeply trained name preferences. This fires inside the actual prompt.
        base += " " + PROMPT_NAME_BAN

        return base

    @staticmethod
    def _get_variation_instruction(index: int) -> str:
        variations = [
            "Add an unexpected twist ending and strong emotional arc.",
            "Use vivid sensory details and keep the pacing brisk.",
            "Write from the perspective of an unreliable narrator.",
            "Focus on a humoristic tone with sharp dialogue.",
            "Make it a philosophical parable with symbolic motifs."
        ]
        return variations[index % len(variations)]

    @staticmethod
    def generate_story_varied(model: str, genre: Optional[str] = None, tone: Optional[str] = None,
                              word_count: int = DEFAULT_WORD_COUNT, custom_prompt: Optional[str] = None,
                              template: Optional[str] = None, variant: int = 0) -> Tuple[str, str]:
        """Generate a storytelling variant for style diversity."""
        base_prompt = StoryGenerator._build_prompt(genre, tone, word_count, custom_prompt, template)
        variation = StoryGenerator._get_variation_instruction(variant)
        prompt = f"{base_prompt} {variation}"
        data = {"model": model, "prompt": prompt, "stream": False}
        try:
            response = requests.post(StoryGenerator.OLLAMA_API_URL, json=data, timeout=300)
            response.raise_for_status()
            result = response.json()
            story_text = result.get("response", "")
            return story_text, model
        except requests.exceptions.RequestException as e:
            raise Exception(f"Error generating varied story with {model}: {e}")

    @staticmethod
    def generate_story_variants(model: str, genre: Optional[str] = None, tone: Optional[str] = None,
                                word_count: int = DEFAULT_WORD_COUNT, custom_prompt: Optional[str] = None,
                                template: Optional[str] = None, variants: int = 3) -> List[Tuple[str, str]]:
        """Generate multiple variant stories for variety mode."""
        stories = []
        for i in range(variants):
            story, model_used = StoryGenerator.generate_story_varied(
                model, genre, tone, word_count, custom_prompt, template, variant=i
            )
            stories.append((story, model_used))
        return stories

    @staticmethod
    def format_story_for_platform(story_text: str, platform: str = 'X', max_chars: int = DEFAULT_MAX_CHARS) -> str:
        """Format story for social platform posting."""
        if not story_text:
            return ""
        text = "\n".join(line.strip() for line in story_text.splitlines() if line.strip())
        if platform.lower() in ['x', 'twitter', 'threads']:
            summary = text[:max_chars-30].rstrip()
            if len(summary) < len(text):
                summary += '...'
            return f"{summary}\n\n#ShortStory #Fiction"
        if platform.lower() == 'facebook':
            summary = text[:max_chars*2].rstrip()
            if len(summary) < len(text):
                summary += '...'
            return f"{summary}\n\nRead more in the app or save for later!"
        if platform.lower() == 'instagram':
            caption = text[:max_chars-20].rstrip()
            if len(caption) < len(text):
                caption += '...'
            return f"{caption}\n\n#fiction #story"
        result = text[:max_chars].rstrip()
        if len(result) < len(text):
            result += '...'
        return result


    @staticmethod
    def chat_streaming(model: str, messages: list, stop_event=None) -> Iterator[str]:
        """
        Multi-turn chat using /api/chat endpoint.
        messages: list of {"role": "user"|"assistant"|"system", "content": str}
        Yields text chunks as they stream in.
        """
        data = {
            "model": model,
            "messages": messages,
            "stream": True,
        }
        try:
            response = requests.post(
                StoryGenerator.CHAT_API_URL,
                json=data, timeout=300, stream=True
            )
            response.raise_for_status()
            for line in response.iter_lines():
                if stop_event and stop_event.is_set():
                    break
                if line:
                    chunk = json.loads(line)
                    text = chunk.get("message", {}).get("content", "")
                    if text:
                        yield text
                    if chunk.get("done", False):
                        break
        except requests.exceptions.RequestException as e:
            raise Exception(f"Chat error with {model}: {e}")

    @staticmethod
    def generate_story_streaming(model: str, genre: Optional[str] = None, tone: Optional[str] = None,
                                word_count: int = DEFAULT_WORD_COUNT, custom_prompt: Optional[str] = None,
                                template: Optional[str] = None, stop_event=None) -> Iterator[str]:
        """Generate a story using streaming - yields text chunks as they arrive."""
        prompt = StoryGenerator._build_prompt(genre, tone, word_count, custom_prompt, template)
        data = {"model": model, "prompt": prompt, "stream": True}
        try:
            response = requests.post(StoryGenerator.OLLAMA_API_URL, json=data, timeout=300, stream=True)
            response.raise_for_status()
            for line in response.iter_lines():
                if stop_event and stop_event.is_set():
                    break
                if line:
                    chunk = json.loads(line)
                    text = chunk.get("response", "")
                    if text:
                        yield text
                    if chunk.get("done", False):
                        break
        except requests.exceptions.RequestException as e:
            raise Exception(f"Error generating story with {model}: {e}")

    @staticmethod
    def generate_story(model: str, genre: Optional[str] = None, tone: Optional[str] = None,
                      word_count: int = DEFAULT_WORD_COUNT, custom_prompt: Optional[str] = None,
                      template: Optional[str] = None) -> Tuple[str, str]:
        """Generate a story using the specified Ollama model (non-streaming)."""
        prompt = StoryGenerator._build_prompt(genre, tone, word_count, custom_prompt, template)
        data = {"model": model, "prompt": prompt, "stream": False}
        try:
            response = requests.post(StoryGenerator.OLLAMA_API_URL, json=data, timeout=300)
            response.raise_for_status()
            result = response.json()
            story_text = result.get("response", "")
            return story_text, model
        except requests.exceptions.RequestException as e:
            raise Exception(f"Error generating story with {model}: {e}")


class StoriesManager:
    """Manages story storage and comparison."""

    RATINGS_FILE = "story_ratings.json"

    def __init__(self, ratings_file: Optional[str] = None, stories_dir: Optional[str] = None) -> None:
        self.STORIES_DIR = stories_dir or os.path.join(os.getcwd(), "generated_stories")
        if not os.path.exists(self.STORIES_DIR):
            os.makedirs(self.STORIES_DIR)
        self.ratings_path = ratings_file or os.path.join(self.STORIES_DIR, self.RATINGS_FILE)
        self._load_ratings()

    def _load_ratings(self) -> None:
        """Load ratings from file or create empty database."""
        if os.path.exists(self.ratings_path):
            try:
                with open(self.ratings_path, 'r', encoding='utf-8', errors='replace') as f:
                    self.ratings = json.load(f)
                for filename, data in self.ratings.items():
                    if not isinstance(data, dict):
                        self.ratings[filename] = {'rating': 0, 'favorite': False}
                        continue
                    rating = data.get('rating', 0)
                    if not isinstance(rating, int) or rating < 0 or rating > 5:
                        data['rating'] = 0
                    if not isinstance(data.get('favorite'), bool):
                        data['favorite'] = False
            except (IOError, json.JSONDecodeError) as e:
                logger.warning(f"Failed to load ratings file: {e}")
                self.ratings = {}
        else:
            self.ratings = {}

    def _save_ratings(self) -> None:
        try:
            with open(self.ratings_path, 'w', encoding='utf-8', errors='replace') as f:
                json.dump(self.ratings, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Error saving ratings: {e}")

    def set_rating(self, filename: str, rating: int) -> bool:
        if rating < 0 or rating > 5:
            return False
        if filename not in self.ratings:
            self.ratings[filename] = {'rating': 0, 'favorite': False}
        self.ratings[filename]['rating'] = rating
        self._save_ratings()
        return True

    def get_rating(self, filename: str) -> int:
        if filename in self.ratings:
            return self.ratings[filename].get('rating', 0)
        return 0

    def set_favorite(self, filename: str, is_favorite: bool) -> bool:
        if filename not in self.ratings:
            self.ratings[filename] = {'rating': 0, 'favorite': False}
        self.ratings[filename]['favorite'] = is_favorite
        self._save_ratings()
        return True

    def get_favorite(self, filename: str) -> bool:
        if filename in self.ratings:
            return self.ratings[filename].get('favorite', False)
        return False

    def get_rating_stats(self) -> dict:
        if not self.ratings:
            return {'average': 0, 'total_rated': 0, 'distribution': {}, 'favorites': 0}
        ratings_list = [r['rating'] for r in self.ratings.values() if r['rating'] > 0]
        favorites_count = sum(1 for r in self.ratings.values() if r['favorite'])
        distribution = {i: 0 for i in range(1, 6)}
        for rating in ratings_list:
            if rating > 0:
                distribution[rating] += 1
        return {
            'average': sum(ratings_list) / len(ratings_list) if ratings_list else 0,
            'total_rated': len(ratings_list),
            'distribution': distribution,
            'favorites': favorites_count
        }

    def get_generation_statistics(self) -> dict:
        all_stories = self.get_all_stories_metadata()
        if not all_stories:
            return {
                'total_stories': 0, 'total_words': 0, 'average_words': 0,
                'models': {}, 'genres': {}, 'tones': {}, 'generation_timeline': {},
                'earliest_date': None, 'latest_date': None
            }
        models: Dict[str, int] = {}
        for story in all_stories:
            model = story['model']
            models[model] = models.get(model, 0) + 1
        genres: Dict[str, int] = {}
        for story in all_stories:
            if story['genre']:
                genres[story['genre']] = genres.get(story['genre'], 0) + 1
        tones: Dict[str, int] = {}
        for story in all_stories:
            if story['tone']:
                tones[story['tone']] = tones.get(story['tone'], 0) + 1
        timeline: Dict[str, int] = {}
        dates_list: List[datetime] = []
        for story in all_stories:
            if story['generated_date']:
                date_key = story['generated_date'].strftime('%Y-%m-%d')
                timeline[date_key] = timeline.get(date_key, 0) + 1
                dates_list.append(story['generated_date'])
        total_words = sum(story['word_count'] for story in all_stories)
        model_performance: Dict[str, Dict[str, Any]] = {}
        for story in all_stories:
            model = story['model']
            filename = story['filename']
            rating = self.get_rating(filename)
            favorite = self.get_favorite(filename)
            if model not in model_performance:
                model_performance[model] = {'count': 0, 'rated_count': 0, 'rating_sum': 0, 'avg_rating': 0, 'favorites': 0}
            mp = model_performance[model]
            mp['count'] += 1
            if rating > 0:
                mp['rated_count'] += 1
                mp['rating_sum'] += rating
            if favorite:
                mp['favorites'] += 1
        for model, mp in model_performance.items():
            mp['avg_rating'] = (mp['rating_sum'] / mp['rated_count']) if mp['rated_count'] else 0
        return {
            'total_stories': len(all_stories),
            'total_words': total_words,
            'average_words': total_words / len(all_stories) if all_stories else 0,
            'models': models, 'genres': genres, 'tones': tones,
            'generation_timeline': timeline,
            'earliest_date': min(dates_list) if dates_list else None,
            'latest_date': max(dates_list) if dates_list else None,
            'rating_stats': self.get_rating_stats(),
            'model_performance': model_performance
        }

    def save_story(self, story_text: str, model: str, genre: Optional[str] = None,
                   tone: Optional[str] = None, filename: Optional[str] = None) -> str:
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"story_{timestamp}.md"
        elif not filename.endswith(".md"):
            filename += ".md"
        filepath = os.path.join(self.STORIES_DIR, filename)
        metadata = f"""---
**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**Model:** {model}
{f'**Genre:** {genre}' if genre else ''}
{f'**Tone:** {tone}' if tone else ''}
---

# Generated Story

{story_text}
"""
        with open(filepath, 'w', encoding='utf-8', errors='replace') as f:
            f.write(metadata)
        return filepath

    def load_story(self, filename: str) -> Optional[str]:
        filepath = os.path.join(self.STORIES_DIR, filename)
        if os.path.exists(filepath):
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        return None

    def _parse_story_lines(self, story_content: str) -> List[str]:
        """Extract story text lines from a saved .md file, stripping frontmatter."""
        lines = story_content.split('\n')
        story_text_lines = []
        in_frontmatter = False
        for line in lines:
            if line.strip() == '---':
                in_frontmatter = not in_frontmatter
                continue
            if not in_frontmatter and line.strip():
                story_text_lines.append(line)
        return story_text_lines

    def _get_export_path(self, filename: str, ext: str, output_path: Optional[str] = None) -> str:
        """Return validated export path, creating the exports dir if needed."""
        if output_path is not None:
            return output_path
        export_dir = os.path.join(self.STORIES_DIR, "exports")
        if not os.path.exists(export_dir):
            os.makedirs(export_dir)
        return os.path.join(export_dir, filename.replace('.md', ext))

    def _load_story_for_export(self, filename: str) -> tuple:
        """Shared prelude for all export methods. Returns (story_text_lines, metadata, content)."""
        story_content = self.load_story(filename)
        if not story_content:
            raise Exception(f"Story '{filename}' not found")
        metadata = self.get_story_metadata(filename)
        if not metadata:
            raise Exception(f"Could not parse metadata for story '{filename}'")
        story_text_lines = self._parse_story_lines(story_content)
        return story_text_lines, metadata, story_content

    def export_to_pdf(self, filename: str, output_path: Optional[str] = None, include_metadata: bool = True) -> str:
        story_text_lines, metadata, _ = self._load_story_for_export(filename)
        output_path = self._get_export_path(filename, '.pdf', output_path)
        doc = SimpleDocTemplate(output_path, pagesize=letter,
                               rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        story_elements = []
        styles = getSampleStyleSheet()
        if include_metadata:
            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'],
                fontSize=24, textColor=colors.HexColor('#2196F3'), spaceAfter=6, alignment=1)
            if metadata['title']:
                story_elements.append(Paragraph(metadata['title'], title_style))
            metadata_data = []
            if metadata['generated_date']:
                metadata_data.append(['Date:', metadata['generated_date'].strftime('%Y-%m-%d %H:%M:%S')])
            metadata_data.append(['Model:', metadata['model']])
            if metadata['genre']:
                metadata_data.append(['Genre:', metadata['genre']])
            if metadata['tone']:
                metadata_data.append(['Tone:', metadata['tone']])
            metadata_data.append(['Words:', str(metadata['word_count'])])
            if metadata_data:
                meta_table = Table(metadata_data, colWidths=[1.5 * inch, 4 * inch])
                meta_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#666666')),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ]))
                story_elements.append(meta_table)
            story_elements.append(Spacer(1, 0.3 * inch))
            story_elements.append(Paragraph("_" * 80, styles['Normal']))
            story_elements.append(Spacer(1, 0.2 * inch))
        for line in story_text_lines:
            if line.startswith('# '):
                story_elements.append(Paragraph(line[2:], styles['Heading2']))
            elif line.strip():
                story_elements.append(Paragraph(line, styles['Normal']))
            else:
                story_elements.append(Spacer(1, 0.1 * inch))
        doc.build(story_elements)
        return output_path

    def export_to_docx(self, filename: str, output_path: Optional[str] = None, include_metadata: bool = True) -> str:
        story_text_lines, metadata, _ = self._load_story_for_export(filename)
        output_path = self._get_export_path(filename, '.docx', output_path)
        doc = Document()
        if include_metadata and metadata['title']:
            title = doc.add_heading(metadata['title'], level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if title.runs:
                title.runs[0].font.color.rgb = RGBColor(33, 150, 243)
        if include_metadata:
            metadata_table = doc.add_table(rows=0, cols=2)
            metadata_table.style = 'Light Grid Accent 1'
            if metadata['generated_date']:
                row = metadata_table.add_row().cells
                row[0].text = 'Date'; row[1].text = metadata['generated_date'].strftime('%Y-%m-%d %H:%M:%S')
            row = metadata_table.add_row().cells
            row[0].text = 'Model'; row[1].text = metadata['model']
            if metadata['genre']:
                row = metadata_table.add_row().cells
                row[0].text = 'Genre'; row[1].text = metadata['genre']
            if metadata['tone']:
                row = metadata_table.add_row().cells
                row[0].text = 'Tone'; row[1].text = metadata['tone']
            row = metadata_table.add_row().cells
            row[0].text = 'Word Count'; row[1].text = str(metadata['word_count'])
            doc.add_paragraph()
        for line in story_text_lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=2)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=3)
            elif line.strip():
                doc.add_paragraph(line)
            else:
                doc.add_paragraph()
        doc.save(output_path)
        return output_path

    def export_to_txt(self, filename: str, output_path: Optional[str] = None, include_metadata: bool = True) -> str:
        story_text_lines, metadata, _ = self._load_story_for_export(filename)
        output_path = self._get_export_path(filename, '.txt', output_path)
        output_text = []
        if include_metadata:
            if metadata['title']:
                output_text.append(metadata['title'])
                output_text.append('=' * len(metadata['title']))
            output_text.append('')
            if metadata['generated_date']:
                output_text.append(f"Generated: {metadata['generated_date'].strftime('%Y-%m-%d %H:%M:%S')}")
            output_text.append(f"Model: {metadata['model']}")
            if metadata['genre']: output_text.append(f"Genre: {metadata['genre']}")
            if metadata['tone']:  output_text.append(f"Tone: {metadata['tone']}")
            output_text.append(f"Word Count: {metadata['word_count']}")
            output_text.append(''); output_text.append('-' * 80); output_text.append('')
        for line in story_text_lines:
            if line.startswith('# '):
                output_text.append('\n' + line[2:].upper())
            elif line.startswith('## '):
                output_text.append('\n' + line[3:])
            else:
                output_text.append(line)
        with open(output_path, 'w', encoding='utf-8', errors='replace') as f:
            f.write('\n'.join(output_text))
        return output_path


    def get_saved_stories(self) -> List[str]:
        """Get list of saved story filenames."""
        if os.path.exists(self.STORIES_DIR):
            return [f for f in os.listdir(self.STORIES_DIR) if f.endswith('.md')]
        return []

    def get_story_metadata(self, filename: str) -> Optional[Dict[str, Any]]:
        """Extract metadata from a saved story file."""
        filepath = os.path.join(self.STORIES_DIR, filename)
        if not os.path.exists(filepath):
            return None
        metadata: Dict[str, Any] = {
            'filename': filename, 'filepath': filepath, 'title': 'Unknown',
            'generated_date': None, 'model': 'Unknown', 'genre': None,
            'tone': None, 'word_count': 0
        }
        try:
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
            lines = content.split('\n')
            in_frontmatter = False
            story_content = ""
            for line in lines:
                if line.strip() == '---':
                    in_frontmatter = not in_frontmatter
                    continue
                if in_frontmatter:
                    if line.startswith('**Generated:**'):
                        try:
                            date_str = line.replace('**Generated:**', '').strip()
                            metadata['generated_date'] = datetime.strptime(date_str, '%Y-%m-%d %H:%M:%S')
                        except ValueError:
                            pass
                    elif line.startswith('**Model:**'):
                        metadata['model'] = line.replace('**Model:**', '').strip()
                    elif line.startswith('**Genre:**'):
                        metadata['genre'] = line.replace('**Genre:**', '').strip()
                    elif line.startswith('**Tone:**'):
                        metadata['tone'] = line.replace('**Tone:**', '').strip()
                else:
                    story_content += line + '\n'
            # Extract title — skip generic headers, find real ## title or use filename
            _GENERIC = {'Generated Story', 'Unknown', ''}
            if story_content.strip():
                _sc_lines = story_content.strip().split('\n')
                first_line = _sc_lines[0]
                if first_line.startswith('# '):
                    metadata['title'] = first_line[2:].strip()
                else:
                    metadata['title'] = first_line[:50] + '...' if len(first_line) > 50 else first_line
                # If H1 is generic, look for a ## subheading (the real story title)
                if metadata['title'] in _GENERIC or not metadata['title'].strip():
                    for _ln in _sc_lines[1:]:
                        _s = _ln.strip()
                        if _s.startswith('## '):
                            metadata['title'] = _s[3:].strip(); break
                        elif _s.startswith('# ') and _s[2:].strip() not in _GENERIC:
                            metadata['title'] = _s[2:].strip(); break
            # Last resort: filename without extension
            if metadata['title'] in _GENERIC or not metadata['title'].strip():
                metadata['title'] = os.path.splitext(filename)[0]
            metadata['word_count'] = len(story_content.split())
        except Exception as e:
            logger.warning(f"Error reading metadata for {filename}: {e}")
        metadata['rating'] = self.get_rating(filename)
        metadata['favorite'] = self.get_favorite(filename)
        return metadata

    def get_all_stories_metadata(self) -> List[dict]:
        """Get metadata for all saved stories."""
        stories = []
        for filename in self.get_saved_stories():
            metadata = self.get_story_metadata(filename)
            if metadata:
                stories.append(metadata)
        return stories

    def filter_stories(self, search_term: Optional[str] = None, genre: Optional[str] = None,
                      model: Optional[str] = None, date_from: Optional[datetime] = None,
                      date_to: Optional[datetime] = None, min_words: Optional[int] = None,
                      max_words: Optional[int] = None, min_rating: Optional[int] = None,
                      favorites_only: bool = False, sort_by: str = 'date_desc') -> List[dict]:
        """Filter and sort stories based on criteria."""
        stories = self.get_all_stories_metadata()
        if genre:
            stories = [s for s in stories if s['genre'] and s['genre'].lower() == genre.lower()]
        if model:
            stories = [s for s in stories if s['model'] and s['model'].lower() == model.lower()]
        if min_words is not None:
            stories = [s for s in stories if s['word_count'] >= min_words]
        if max_words is not None:
            stories = [s for s in stories if s['word_count'] <= max_words]
        if min_rating is not None:
            stories = [s for s in stories if s['rating'] >= min_rating]
        if favorites_only:
            stories = [s for s in stories if s['favorite']]
        if search_term:
            term = search_term.lower()
            filtered = []
            for story in stories:
                if term in story['title'].lower() or term in story['model'].lower():
                    filtered.append(story)
                    continue
                try:
                    content = self.load_story(story['filename']) or ''
                    if term in content.lower():
                        filtered.append(story)
                except Exception:
                    pass
            stories = filtered
        sort_map = {
            'date_desc': lambda s: s['generated_date'] or datetime.min,
            'date_asc':  lambda s: s['generated_date'] or datetime.min,
            'words_desc': lambda s: s['word_count'],
            'words_asc':  lambda s: s['word_count'],
            'model':      lambda s: s['model'].lower(),
            'rating_desc': lambda s: s['rating'],
        }
        reverse_map = {
            'date_desc': True, 'date_asc': False,
            'words_desc': True, 'words_asc': False,
            'model': False, 'rating_desc': True,
        }
        key_fn = sort_map.get(sort_by, sort_map['date_desc'])
        reverse = reverse_map.get(sort_by, True)
        stories.sort(key=key_fn, reverse=reverse)
        return stories


class StoryMetrics:
    """Calculate quality metrics for generated stories."""

    @staticmethod
    def calculate_metrics(story_text: str) -> dict:
        if not story_text:
            return {'word_count': 0, 'sentence_count': 0, 'readability_score': 0,
                    'word_variety': 0, 'complex_word_ratio': 0, 'dialogue_ratio': 0,
                    'sentence_variety': 0}
        words = story_text.split()
        word_count = len(words)
        sentences = re.split(r'[.!?]+', story_text)
        sentences = [s.strip() for s in sentences if s.strip()]
        sentence_count = len(sentences)
        unique_words = len(set(w.lower().strip('.,!?";:') for w in words))
        word_variety = round((unique_words / word_count * 100), 1) if word_count > 0 else 0
        complex_words = [w for w in words if len(w) > 8]
        complex_word_ratio = round((len(complex_words) / word_count * 100), 1) if word_count > 0 else 0
        dialogue_lines = [l for l in story_text.split('\n') if '"' in l or "'" in l]
        dialogue_ratio = round((len(dialogue_lines) / max(len(story_text.split('\n')), 1) * 100), 1)
        if sentence_count > 1:
            sentence_lengths = [len(s.split()) for s in sentences]
            avg_len = sum(sentence_lengths) / len(sentence_lengths)
            variance = sum((l - avg_len) ** 2 for l in sentence_lengths) / len(sentence_lengths)
            sentence_variety = round(min(variance ** 0.5 / max(avg_len, 1) * 10, 10), 1)
        else:
            sentence_variety = 0
        avg_words_per_sentence = word_count / max(sentence_count, 1)
        readability_score = round(max(0, min(10, 10 - (avg_words_per_sentence - 15) * 0.2)), 1)
        return {
            'word_count': word_count,
            'sentence_count': sentence_count,
            'readability_score': readability_score,
            'word_variety': word_variety,
            'complex_word_ratio': complex_word_ratio,
            'dialogue_ratio': dialogue_ratio,
            'sentence_variety': sentence_variety
        }


class PresetsManager:
    """Manages story generation presets."""

    PRESETS_FILE = "presets.json"

    def __init__(self, presets_dir: Optional[str] = None) -> None:
        self.presets_dir = presets_dir or os.path.join(os.getcwd(), "generated_stories")
        self.presets_path = os.path.join(self.presets_dir, self.PRESETS_FILE)
        self._load_presets()

    def _load_presets(self) -> None:
        if os.path.exists(self.presets_path):
            try:
                with open(self.presets_path, 'r', encoding='utf-8') as f:
                    self.presets = json.load(f)
            except (IOError, json.JSONDecodeError):
                self.presets = []
        else:
            self.presets = []

    def _save_presets(self) -> None:
        try:
            with open(self.presets_path, 'w', encoding='utf-8') as f:
                json.dump(self.presets, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Error saving presets: {e}")

    def get_presets(self) -> List[dict]:
        return self.presets.copy()

    def get_preset(self, name: str) -> Optional[dict]:
        return next((p for p in self.presets if p['name'] == name), None)

    def save_preset(self, name: str, genre: Optional[str] = None, tone: Optional[str] = None,
                   word_count: int = DEFAULT_WORD_COUNT, custom_prompt: Optional[str] = None) -> bool:
        if self.get_preset(name):
            return False
        preset = {'name': name, 'genre': genre, 'tone': tone,
                  'word_count': word_count, 'custom_prompt': custom_prompt,
                  'created_at': datetime.now().isoformat()}
        self.presets.append(preset)
        self._save_presets()
        return True

    def update_preset(self, name: str, genre: Optional[str] = None, tone: Optional[str] = None,
                     word_count: int = DEFAULT_WORD_COUNT, custom_prompt: Optional[str] = None) -> bool:
        preset = self.get_preset(name)
        if not preset:
            return False
        preset.update({'genre': genre, 'tone': tone, 'word_count': word_count,
                       'custom_prompt': custom_prompt, 'updated_at': datetime.now().isoformat()})
        self._save_presets()
        return True

    def delete_preset(self, name: str) -> bool:
        preset = self.get_preset(name)
        if not preset:
            return False
        self.presets.remove(preset)
        self._save_presets()
        return True


class ContentQueueManager:
    """Manages a queue of scheduled content posts."""

    QUEUE_FILE = "story_queue.json"

    def __init__(self, queue_dir: Optional[str] = None) -> None:
        self.queue_dir = queue_dir or os.path.join(os.getcwd(), "generated_stories")
        self.queue_path = os.path.join(self.queue_dir, self.QUEUE_FILE)
        self._load_queue()

    def _load_queue(self) -> None:
        if os.path.exists(self.queue_path):
            try:
                with open(self.queue_path, 'r', encoding='utf-8') as f:
                    raw = json.load(f)
                self.queue = []
                for item in raw:
                    if 'scheduled_time' in item and isinstance(item['scheduled_time'], str):
                        try:
                            item['scheduled_time'] = datetime.fromisoformat(item['scheduled_time'])
                        except ValueError:
                            pass
                    self.queue.append(item)
            except (IOError, json.JSONDecodeError):
                self.queue = []
        else:
            self.queue = []

    def _save_queue(self) -> None:
        try:
            serializable = []
            for item in self.queue:
                item_copy = item.copy()
                if isinstance(item_copy.get('scheduled_time'), datetime):
                    item_copy['scheduled_time'] = item_copy['scheduled_time'].isoformat()
                serializable.append(item_copy)
            with open(self.queue_path, 'w', encoding='utf-8') as f:
                json.dump(serializable, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Error saving queue: {e}")

    def get_queue(self) -> List[dict]:
        return self.queue.copy()

    def add_to_queue(self, item: dict) -> bool:
        if 'id' not in item:
            item['id'] = f"item_{datetime.now().timestamp()}"
        self.queue.append(item)
        self._save_queue()
        return True

    def remove_from_queue(self, item_id: str) -> bool:
        item = next((i for i in self.queue if i.get('id') == item_id), None)
        if not item:
            return False
        self.queue.remove(item)
        self._save_queue()
        return True

    def update_queue_item(self, item_id: str, updates: dict) -> bool:
        item = next((i for i in self.queue if i.get('id') == item_id), None)
        if not item:
            return False
        item.update(updates)
        self._save_queue()
        return True

    def get_scheduled_items(self, before_time: datetime) -> List[dict]:
        return [i for i in self.queue
                if isinstance(i.get('scheduled_time'), datetime)
                and i['scheduled_time'] <= before_time
                and i.get('status') not in ('executed', 'failed')]

    def execute_queue_item(self, item_id: str, stories_manager: 'StoriesManager') -> bool:
        item = next((i for i in self.queue if i.get('id') == item_id), None)
        if not item:
            return False
        try:
            item['status'] = 'executed'
            item['executed_at'] = datetime.now().isoformat()
            logger.info(f"Executed queue item: {item.get('story_title', item_id)}")
            self._save_queue()
            return True
        except Exception as e:
            item['status'] = 'failed'
            item['error'] = str(e)
            self._save_queue()
            logger.error(f"Failed to execute queue item {item_id}: {e}")
            return False

    def get_performance_stats(self) -> dict:
        return {
            'total_posts': len(self.queue),
            'total_views': 0, 'total_likes': 0,
            'total_shares': 0, 'total_comments': 0,
            'avg_engagement': 0.0
        }

    @staticmethod
    def format_story_as_thread(story_text: str, max_chars: int = DEFAULT_MAX_CHARS) -> List[str]:
        if not story_text:
            return []
        paragraphs = [p.strip() for p in story_text.split('\n\n') if p.strip()]
        segments = []
        current = ""
        for para in paragraphs:
            if len(current) + len(para) + 2 <= max_chars:
                current = (current + "\n\n" + para).strip()
            else:
                if current:
                    segments.append(current)
                if len(para) <= max_chars:
                    current = para
                else:
                    words = para.split()
                    current = ""
                    for word in words:
                        if len(current) + len(word) + 1 <= max_chars:
                            current = (current + " " + word).strip()
                        else:
                            if current:
                                segments.append(current)
                            current = word
        if current:
            segments.append(current)
        return segments

class StoryBible:
    """
    Living document for a long-form novel.
    Tracks characters, places, plot threads, timeline, and voice —
    injected into every chapter generation so the model stays consistent
    across 200k-300k words without needing the full text in context.
    """

    VERSION = "1.0"

    def __init__(self, novel_dir: str, title: str = "Untitled Novel") -> None:
        self.novel_dir  = novel_dir
        self.title      = title
        self.bible_path = os.path.join(novel_dir, "bible.json")
        self._data: dict = self._empty()
        if os.path.exists(self.bible_path):
            self._load()

    # ── Schema ────────────────────────────────────────────────────────────────
    def _empty(self) -> dict:
        return {
            "version":       self.VERSION,
            "title":         self.title,
            "genre":         "",
            "tone":          "",
            "premise":       "",
            "characters":    {},   # name -> {description, role, arc, last_seen}
            "places":        {},   # name -> {description, significance}
            "rules":         [],   # world rules / magic system / tech
            "timeline":      [],   # [{"chapter": n, "event": str}]
            "open_threads":  [],   # unresolved plot threads
            "closed_threads":[],   # resolved threads
            "voice_notes":   "",   # style/tone reminders
            "chapter_count": 0,
            "last_excerpt":  "",   # last ~400 words of most recent chapter
        }

    def _load(self) -> None:
        try:
            with open(self.bible_path, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
            self._data.update(loaded)
        except Exception as e:
            logger.warning(f"Could not load bible: {e}")

    def save(self) -> None:
        os.makedirs(self.novel_dir, exist_ok=True)
        with open(self.bible_path, 'w', encoding='utf-8') as f:
            json.dump(self._data, f, indent=2, ensure_ascii=False)

    # ── Accessors ─────────────────────────────────────────────────────────────
    @property
    def data(self) -> dict: return self._data

    def get(self, key: str, default: Any = None) -> Any:
        return self._data.get(key, default)

    def set(self, key: str, value: Any) -> None:
        self._data[key] = value
        self.save()

    def update(self, updates: dict) -> None:
        self._data.update(updates)
        self.save()

    def add_character(self, name: str, description: str, role: str = "", arc: str = "") -> None:
        self._data["characters"][name] = {
            "description": description,
            "role": role,
            "arc": arc,
            "last_seen": self._data["chapter_count"]
        }
        self.save()

    def add_place(self, name: str, description: str, significance: str = "") -> None:
        self._data["places"][name] = {
            "description": description,
            "significance": significance
        }
        self.save()

    def add_timeline_event(self, chapter: int, event: str) -> None:
        self._data["timeline"].append({"chapter": chapter, "event": event})
        self.save()

    def add_open_thread(self, thread: str) -> None:
        if thread not in self._data["open_threads"]:
            self._data["open_threads"].append(thread)
            self.save()

    def close_thread(self, thread: str) -> None:
        if thread in self._data["open_threads"]:
            self._data["open_threads"].remove(thread)
            self._data["closed_threads"].append(thread)
            self.save()

    def set_last_excerpt(self, text: str, max_words: int = 350) -> None:
        """Store the last N words of the latest chapter for continuity."""
        words = text.split()
        excerpt = " ".join(words[-max_words:]) if len(words) > max_words else text
        self._data["last_excerpt"] = excerpt
        self.save()

    # ── Render to context ─────────────────────────────────────────────────────
    def render_for_context(self, max_chars: int = 2500) -> str:
        """
        Render a compact version of the bible for injection into prompts.
        Stays under max_chars to leave room for actual chapter content.
        """
        d = self._data
        lines = [f"=== STORY BIBLE: {d['title']} ==="]

        if d.get("premise"):
            lines.append(f"PREMISE: {d['premise'][:300]}")
        if d.get("genre") or d.get("tone"):
            lines.append(f"GENRE/TONE: {d.get('genre','')} / {d.get('tone','')}")

        if d["characters"]:
            lines.append("\nCHARACTERS:")
            for name, info in d["characters"].items():
                desc = info.get("description","")[:120]
                arc  = info.get("arc","")[:80]
                lines.append(f"  {name}: {desc}" + (f" | Arc: {arc}" if arc else ""))

        if d["places"]:
            lines.append("\nPLACES:")
            for name, info in list(d["places"].items())[:8]:
                lines.append(f"  {name}: {info.get('description','')[:100]}")

        if d["rules"]:
            lines.append("\nWORLD RULES:")
            for rule in d["rules"][:6]:
                lines.append(f"  - {rule[:120]}")

        if d["open_threads"]:
            lines.append("\nOPEN PLOT THREADS:")
            for t in d["open_threads"][:8]:
                lines.append(f"  - {t[:120]}")

        if d["timeline"]:
            lines.append("\nRECENT TIMELINE:")
            for entry in d["timeline"][-6:]:
                lines.append(f"  Ch{entry['chapter']}: {entry['event'][:100]}")

        if d.get("voice_notes"):
            lines.append(f"\nVOICE: {d['voice_notes'][:200]}")

        text = "\n".join(lines)

        # Trim to fit
        if len(text) > max_chars:
            text = text[:max_chars] + "\n[bible trimmed]"

        return text


class NovelManager:
    """
    Manages a long-form novel: chapters, assembly, bible updates.
    Novels live in generated_stories/novels/{slug}/
    """

    def __init__(self, stories_dir: str) -> None:
        self.novels_dir = os.path.join(stories_dir, "novels")
        os.makedirs(self.novels_dir, exist_ok=True)

    def list_novels(self) -> list:
        if not os.path.exists(self.novels_dir):
            return []
        novels = []
        for slug in os.listdir(self.novels_dir):
            novel_dir = os.path.join(self.novels_dir, slug)
            bible_path = os.path.join(novel_dir, "bible.json")
            if os.path.isdir(novel_dir) and os.path.exists(bible_path):
                try:
                    with open(bible_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    chapters = self._count_chapters(novel_dir)
                    novels.append({
                        "slug":          slug,
                        "title":         data.get("title", slug),
                        "genre":         data.get("genre", ""),
                        "chapter_count": chapters,
                        "novel_dir":     novel_dir,
                    })
                except Exception:
                    pass
        return novels

    def create_novel(self, title: str, genre: str = "", tone: str = "",
                     premise: str = "") -> "StoryBible":
        slug = self._slugify(title)
        novel_dir = os.path.join(self.novels_dir, slug)
        os.makedirs(novel_dir, exist_ok=True)
        bible = StoryBible(novel_dir, title)
        bible.update({"genre": genre, "tone": tone, "premise": premise})
        return bible

    def load_novel(self, slug: str) -> "StoryBible":
        novel_dir = os.path.join(self.novels_dir, slug)
        bible_path = os.path.join(novel_dir, "bible.json")
        if not os.path.exists(bible_path):
            raise FileNotFoundError(f"Novel '{slug}' not found")
        with open(bible_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        bible = StoryBible(novel_dir, data.get("title", slug))
        return bible

    def get_chapter_path(self, bible: "StoryBible", chapter_num: int) -> str:
        return os.path.join(bible.novel_dir, f"chapter_{chapter_num:03d}.md")

    def save_chapter(self, bible: "StoryBible", chapter_num: int,
                     text: str, title: str = "") -> str:
        path = self.get_chapter_path(bible, chapter_num)
        header = f"# Chapter {chapter_num}" + (f": {title}" if title else "") + "\n\n"
        with open(path, 'w', encoding='utf-8') as f:
            f.write(header + text)
        # Update bible
        bible.set_last_excerpt(text)
        bible.set("chapter_count", chapter_num)
        return path

    def load_chapter(self, bible: "StoryBible", chapter_num: int) -> str:
        path = self.get_chapter_path(bible, chapter_num)
        if not os.path.exists(path):
            return ""
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()

    def list_chapters(self, bible: "StoryBible") -> list:
        chapters = []
        for fname in sorted(os.listdir(bible.novel_dir)):
            if fname.startswith("chapter_") and fname.endswith(".md"):
                num_str = fname.replace("chapter_", "").replace(".md", "")
                try:
                    num = int(num_str)
                    path = os.path.join(bible.novel_dir, fname)
                    with open(path, 'r', encoding='utf-8') as f:
                        first_line = f.readline().strip()
                    title = first_line.lstrip("# ").strip() if first_line else f"Chapter {num}"
                    with open(path, 'r', encoding='utf-8') as wc_f:
                        word_count = len(wc_f.read().split())
                    chapters.append({"num": num, "title": title, "word_count": word_count, "path": path})
                except Exception:
                    pass
        return chapters

    def assemble_novel(self, bible: "StoryBible") -> str:
        """Concatenate all chapters into a single manuscript."""
        chapters = self.list_chapters(bible)
        parts = [f"# {bible.data['title']}\n\n"]
        if bible.data.get("premise"):
            parts.append(f"*{bible.data['premise']}*\n\n")
        parts.append("---\n\n")
        for ch in chapters:
            with open(ch["path"], 'r', encoding='utf-8') as f:
                parts.append(f.read() + "\n\n---\n\n")
        return "".join(parts)

    def export_novel_txt(self, bible: "StoryBible") -> str:
        manuscript = self.assemble_novel(bible)
        out_path = os.path.join(bible.novel_dir, "full_novel.txt")
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(manuscript)
        return out_path

    def _count_chapters(self, novel_dir: str) -> int:
        return len([f for f in os.listdir(novel_dir)
                    if f.startswith("chapter_") and f.endswith(".md")])

    def _slugify(self, title: str) -> str:
        slug = re.sub(r'[^\w\s-]', '', title.lower())
        slug = re.sub(r'[\s]+', '_', slug.strip())
        return slug[:50] or "novel"

    # ── Chapter generation ────────────────────────────────────────────────────
    @staticmethod
    def build_chapter_prompt(bible: "StoryBible", chapter_num: int,
                              chapter_brief: str, word_count: int = 2000,
                              custom_instruction: str = "") -> str:
        """
        Build the full prompt for chapter generation.
        Injects the Story Bible context so the model stays consistent.
        """
        bible_context = bible.render_for_context(max_chars=2200)
        last_excerpt  = bible.get("last_excerpt", "")

        parts = [bible_context, ""]

        if last_excerpt and chapter_num > 1:
            parts.append("=== END OF PREVIOUS CHAPTER (last ~350 words) ===")
            parts.append(last_excerpt)
            parts.append("")

        parts.append(f"=== NOW WRITE: Chapter {chapter_num} ===")
        parts.append(f"Chapter brief: {chapter_brief}")
        parts.append("")
        parts.append(
            f"Write Chapter {chapter_num} of '{bible.data['title']}'. "
            f"Target: approximately {word_count} words. "
            f"Begin directly with the chapter prose — no preamble, no summary, no notes. "
            f"Maintain complete consistency with all characters, places, and events "
            f"established in the Story Bible above. "
            f"The chapter should end at a natural stopping point that creates momentum "
            f"for the next chapter."
        )

        # Name ban — injected at prompt level, same as short story generator
        parts.append("\n" + PROMPT_NAME_BAN)

        if custom_instruction:
            parts.append(f"Additional instruction: {custom_instruction}")

        return "\n".join(parts)

    @staticmethod
    def build_bible_update_prompt(bible: "StoryBible", chapter_num: int,
                                   chapter_text: str) -> str:
        """
        Build prompt asking model to extract new facts from a finished chapter
        and return structured JSON for updating the Story Bible.
        """
        return f"""You have just written Chapter {chapter_num} of '{bible.data['title']}'.

Read the chapter below and extract any NEW information that should be added to the Story Bible.
Return ONLY valid JSON — no prose, no markdown fences, just the raw JSON object.

Format:
{{
  "new_characters": [{{"name": "...", "description": "...", "role": "...", "arc": "..."}}],
  "new_places": [{{"name": "...", "description": "...", "significance": "..."}}],
  "new_rules": ["..."],
  "timeline_events": ["brief event description"],
  "new_open_threads": ["..."],
  "closed_threads": ["..."],
  "voice_notes": "any style/tone observations worth remembering (or empty string)"
}}

Only include entries for things that are genuinely NEW — not already in the bible.
If there is nothing new in a category, use an empty list or empty string.

CHAPTER {chapter_num}:
{chapter_text[:4000]}"""

    @staticmethod
    def apply_bible_update(bible: "StoryBible", chapter_num: int,
                            json_str: str) -> list:
        """
        Parse model's JSON bible update and apply it.
        Returns list of what was added.
        """
        added = []

        # Strip markdown fences if model ignored instructions
        clean = re.sub(r'```(?:json)?\s*', '', json_str).strip().rstrip('`').strip()
        # Find first { ... } block
        m = re.search(r'\{.*\}', clean, re.DOTALL)
        if not m:
            logger.warning("No JSON found in bible update response")
            return added

        try:
            data = json.loads(m.group())
        except json.JSONDecodeError as e:
            logger.warning(f"Bible update JSON parse error: {e}")
            return added

        for ch in data.get("new_characters", []):
            if isinstance(ch, dict) and ch.get("name"):
                bible.add_character(ch["name"], ch.get("description",""),
                                     ch.get("role",""), ch.get("arc",""))
                added.append(f"Character: {ch['name']}")

        for pl in data.get("new_places", []):
            if isinstance(pl, dict) and pl.get("name"):
                bible.add_place(pl["name"], pl.get("description",""),
                                 pl.get("significance",""))
                added.append(f"Place: {pl['name']}")

        for rule in data.get("new_rules", []):
            if rule and rule not in bible.data["rules"]:
                bible.data["rules"].append(rule)
                added.append(f"Rule: {rule[:60]}")

        for ev in data.get("timeline_events", []):
            if ev:
                bible.add_timeline_event(chapter_num, ev)
                added.append(f"Timeline: {ev[:60]}")

        for t in data.get("new_open_threads", []):
            if t:
                bible.add_open_thread(t)
                added.append(f"Thread: {t[:60]}")

        for t in data.get("closed_threads", []):
            if t:
                bible.close_thread(t)
                added.append(f"Closed: {t[:60]}")

        if data.get("voice_notes"):
            existing = bible.get("voice_notes", "")
            new_note = data["voice_notes"]
            if new_note and new_note not in existing:
                bible.set("voice_notes", (existing + " " + new_note).strip())
                added.append("Voice notes updated")

        if added:
            bible.save()

        return added
